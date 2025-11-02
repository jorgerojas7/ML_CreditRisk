"""
Genera un documento Word (DOCX) en español con la explicación del notebook
principal y los módulos del paquete ml_creditrisk.

Requisitos:
- python-docx

Salida:
- docs/Explicacion_Notebook_y_Modulos.docx
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullets(document: Document, items: list[str]):
    for it in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)


def build_document() -> Document:
    doc = Document()

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ML Credit Risk Analysis")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Explicación del Notebook y Módulos")
    run2.font.size = Pt(14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    run3.font.size = Pt(10)

    doc.add_page_break()

    # Introducción
    add_heading(doc, "Introducción", level=1)
    add_paragraph(
        doc,
        (
            "Este documento resume el flujo de trabajo del proyecto de análisis de riesgo "
            "crediticio. El enfoque principal es el notebook '02_Feature_Engineering_Modelado.ipynb', "
            "que orquesta: carga de datos, agrupación de variables, construcción de un preprocesador "
            "a partir de 'df_groups_final', cálculo de importancia de variables, filtrado por umbral, "
            "entrenamiento de modelos, predicciones y visualización de resultados.\n\n"
            "Además, se describen los módulos del paquete 'ml_creditrisk', que encapsulan la lógica "
            "reutilizable para facilitar la trazabilidad y la mantenibilidad."
        ),
    )

    # Notebook
    add_heading(doc, "Notebook principal: 02_Feature_Engineering_Modelado.ipynb", level=1)
    add_paragraph(
        doc,
        (
            "El notebook implementa el flujo end-to-end (E2E). Las secciones clave son:\n"
        ),
        bold=True,
    )
    add_bullets(
        doc,
        [
            "Carga de datos (PAKDD2010_*.txt / .XLS).",
            "Agrupación de variables y exclusiones; 'df_groups_final' como fuente única de verdad.",
            "Construcción del preprocesador a partir de los grupos (ColumnTransformer/Pipeline).",
            "Cálculo de importancia con XGBoost y agregación a variables originales.",
            "Filtrado por umbral de importancia → 'preprocessor_filtered' y tabla de variables eliminadas.",
            "Entrenamiento y evaluación de modelos (RF, XGBoost, LightGBM, CatBoost).",
            "Predicciones sobre 'Prediction_Data.txt' con columnas 'score_*'.",
            "Histogramas de score por modelo.",
            "(Opcional) Búsqueda de hiperparámetros con RandomizedSearchCV (flag HPO_ENABLED).",
            "Guardado del preprocesador como artefacto .joblib + metadatos .json.",
        ],
    )

    # Módulos del paquete
    add_heading(doc, "Paquete: ml_creditrisk", level=1)

    add_heading(doc, "feature_grouping.py", level=2)
    add_bullets(
        doc,
        [
            "compute_missing_pct: calcula % de faltantes por columna.",
            "build_groups_raw: genera la tabla de grupos inicial.",
            "make_rectangular: normaliza la tabla de grupos.",
            "parse_df_groups_table: parsea/valida la tabla de grupos mostrada.",
        ],
    )

    add_heading(doc, "preprocessing.py", level=2)
    add_bullets(
        doc,
        [
            "QuantileDiscretizer: discretizador robusto por cuantiles (AGE, MONTHS_IN_THE_JOB).",
            "build_preprocessor_from_groups: arma ColumnTransformer desde 'df_groups_final'.",
            "resumen_columnas: resume columnas de salida del preprocesador.",
        ],
    )

    add_heading(doc, "importance.py", level=2)
    add_bullets(
        doc,
        [
            "build_output_to_raw_mapping: mapea columnas output a variables originales.",
            "train_xgb_and_agg_importances: entrena XGBoost y agrega importancias por variable raw.",
            "build_filtered_preprocessor: crea un preprocesador filtrado por umbral de importancia.",
            "plot_importances: gráfico de importancias.",
            "dropped_variables_table: tabla de variables eliminadas.",
        ],
    )

    add_heading(doc, "models.py", level=2)
    add_bullets(
        doc,
        [
            "get_base_models: registra modelos base (RF, XGBoost, LightGBM, CatBoost).",
            "LeafIndexEncoder / get_xgb_leaves_lr: pipeline de 'GB leaves → OneHot → LR' (opcional).",
            "evaluate_models: evaluación y métricas.",
        ],
    )

    # Artefactos y versionado
    add_heading(doc, "Artefactos y versionado", level=1)
    add_bullets(
        doc,
        [
            "El preprocesador se guarda como 'models/preprocessor_active_<timestamp>.joblib'.",
            "Los metadatos asociados se guardan como '.json' en la misma carpeta.",
            "Por defecto, 'models/*.joblib' y '.json' están ignorados en .gitignore (evitar repos pesados).",
            "Para versionar un artefacto específico, usar 'git add -f' o ajustar .gitignore.",
        ],
    )

    add_paragraph(
        doc,
        (
            "Recomendación: documentar en los metadatos el umbral de importancia utilizado, "
            "los flags relevantes (HPO_ENABLED, USE_TUNED_MODELS) y un hash de columnas de entrada."
        ),
        italic=True,
    )

    return doc


essential_dirs = [
    Path("docs"),
    Path("models"),
]


def main():
    for d in essential_dirs:
        d.mkdir(parents=True, exist_ok=True)

    doc = build_document()
    out_path = Path("docs") / "Explicacion_Notebook_y_Modulos.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    main()
